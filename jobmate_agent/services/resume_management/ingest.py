"""Resume / job 文本文件的通用 ingest 工具函数。

负责完成以下步骤：
1. 校验文件类型与大小。
2. 依据后缀解析 PDF/DOCX/TXT，抽取纯文本。
3. 进行格式归一化，附带警告信息。
4. 将原始文件落盘，并返回结构化结果，供后续持久化到数据库。
"""

from __future__ import annotations

import io
import logging
import os
import re
import unicodedata
import uuid
import mimetypes
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

# 允许的扩展名集合；未来扩展 job description 解析时也可复用。
# ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
# Currently only allow .docx files
ALLOWED_EXTENSIONS = {".docx"}
# Max size limit for uploads (MB). Keep default aligned with existing behavior (10MB)
DEFAULT_MAX_SIZE_MB = int(os.getenv("RESUME_MAX_SIZE_MB", 10))
# 默认原始文件存放路径，所有上传会存一个原文件备份。
DEFAULT_UPLOAD_DIR = Path(__file__).resolve().parents[1] / "data" / "uploads"


class ResumeParsingError(Exception):
    """上传文件解析失败时抛出，统一给调用方处理为 4xx/5xx。"""


@dataclass(frozen=True)
class ResumeIngestResult:
    """封装 ingest 成功后的产出。"""

    file_path: Optional[Path]
    parsed_json: Dict[str, Any]
    warnings: Sequence[str]


def prepare_resume_bytes(
    file_storage: FileStorage,
    *,
    max_size_mb: Optional[int] = None,
) -> Tuple[bytes, str, str]:
    """
    Read the incoming upload stream exactly once, enforce a size cap, and
    return immutable bytes plus normalized filename and content type.

    This centralizes risky stream handling so callers (routes/pipelines) don't
    need to `seek()` or read multiple times.

    Returns: (data, normalized_filename, content_type)
    """
    if not file_storage or not file_storage.filename:
        raise ResumeParsingError("No file supplied.")

    size_limit_mb = DEFAULT_MAX_SIZE_MB if max_size_mb is None else max_size_mb
    size_limit_bytes = size_limit_mb * 1024 * 1024

    # Normalize filename and content type up front
    normalized_name = secure_filename(file_storage.filename)
    content_type = (
        file_storage.content_type
        or mimetypes.guess_type(normalized_name)[0]
        or "application/octet-stream"
    )

    # Read the entire stream exactly once
    # Note: For current size limit (<=10MB), in-memory read is acceptable and
    # matches current implementation behavior.
    try:
        if hasattr(file_storage.stream, "seek"):
            try:
                file_storage.stream.seek(0)
            except Exception:
                # If seek fails, continue and attempt to read anyway
                pass
        raw_bytes = file_storage.stream.read()
        if raw_bytes is None:
            raw_bytes = b""
    except Exception as exc:
        raise ResumeParsingError(f"Failed to read upload stream: {exc}")

    if not raw_bytes:
        raise ResumeParsingError("Empty file.")

    if len(raw_bytes) > size_limit_bytes:
        raise ResumeParsingError("File exceeds size limit.")

    detected_ext = _detect_file_type_from_content(raw_bytes)
    declared_ext = _detect_extension(normalized_name)
    if detected_ext and declared_ext and detected_ext != declared_ext:
        # Allow common aliasing (e.g., .docx but detected zip that contains DOCX parts)
        # Existing parse logic already attempts a DOCX validation when zip is detected
        logger.warning(
            "Filename/content type mismatch: name=%s ext=%s detected=%s",
            normalized_name,
            declared_ext,
            detected_ext,
        )

    return raw_bytes, normalized_name, content_type


def parse_resume_file(
    file_storage: FileStorage,
    *,
    upload_dir: Optional[Path] = None,
    max_size_mb: int = 10,
) -> ResumeIngestResult:
    """解析上传的简历或 JD 文件，并将原始文件保存到磁盘。"""
    if not file_storage or not file_storage.filename:
        raise ResumeParsingError("No file supplied.")

    # Determine if local storage should be used; do not create directories unless enabled
    should_store_local = os.getenv("STORE_LOCAL_UPLOAD_COPY", "0") == "1" or (
        upload_dir is not None
    )
    target_upload_dir: Optional[Path] = (
        upload_dir
        if upload_dir is not None
        else (DEFAULT_UPLOAD_DIR if should_store_local else None)
    )

    # 检查扩展名，确保只处理白名单内的文件类型。
    extension = _detect_extension(file_storage.filename)
    logger.debug(
        f"Detected extension from filename '{file_storage.filename}': {extension}"
    )

    if extension not in ALLOWED_EXTENSIONS:
        raise ResumeParsingError(f"Unsupported file type: {extension}")

    # 使用统一的字节解析，避免重复读取流
    data, normalized_name, content_type = prepare_resume_bytes(
        file_storage, max_size_mb=max_size_mb
    )
    return parse_resume_bytes(
        raw_bytes=data,
        filename=normalized_name,
        content_type=content_type,
        upload_dir=target_upload_dir,
        max_size_mb=max_size_mb,
    )


def parse_resume_bytes(
    *,
    raw_bytes: bytes,
    filename: str,
    content_type: str,
    upload_dir: Optional[Path] = None,
    max_size_mb: int = DEFAULT_MAX_SIZE_MB,
) -> ResumeIngestResult:
    """解析上传的简历或 JD 字节内容，并将原始文件保存到磁盘。

    该函数在上层已统一完成一次读取与大小校验，这里再做幂等校验与类型识别，
    以便在不同入口均可安全调用。
    """
    if not raw_bytes:
        raise ResumeParsingError("Empty file.")

    if len(raw_bytes) > max_size_mb * 1024 * 1024:
        raise ResumeParsingError("File exceeds size limit.")

    # Only prepare local directory when local storage is enabled
    should_store_local = os.getenv("STORE_LOCAL_UPLOAD_COPY", "0") == "1" or (
        upload_dir is not None
    )
    upload_path: Optional[Path] = None
    if should_store_local:
        upload_path = upload_dir if upload_dir is not None else DEFAULT_UPLOAD_DIR
        upload_path.mkdir(parents=True, exist_ok=True)

    # 从文件名提取初始扩展名
    extension = _detect_extension(filename)
    logger.debug(f"Detected extension from filename '{filename}': {extension}")

    if extension not in ALLOWED_EXTENSIONS:
        # 允许后续从内容识别覆盖
        logger.debug(
            f"Extension {extension} not in allowlist; will rely on content detection"
        )

    # 根据内容识别文件类型
    actual_extension = _detect_file_type_from_content(raw_bytes)
    logger.debug(f"Detected file type from content: {actual_extension}")

    if actual_extension and actual_extension != extension:
        logger.warning(
            f"File extension mismatch: filename says {extension}, content says {actual_extension}"
        )
        if extension == ".docx" and actual_extension == ".zip":
            logger.debug("Attempting to validate ZIP file as DOCX...")
            try:
                import zipfile

                with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zip_file:
                    if "word/document.xml" in zip_file.namelist():
                        logger.debug(
                            "ZIP file contains DOCX structure, treating as DOCX"
                        )
                        actual_extension = ".docx"
            except Exception as e:
                logger.warning(f"Failed to validate as DOCX: {e}")

        if actual_extension in ALLOWED_EXTENSIONS:
            extension = actual_extension
            logger.debug(f"Using content-detected extension: {extension}")

    pages, warnings = _parse_by_extension(extension, raw_bytes)
    parsed_json = _build_parsed_payload(pages, warnings)

    stored_path: Optional[Path] = None
    if should_store_local:
        # upload_path must be prepared when should_store_local is True
        assert upload_path is not None
        stored_path = _store_file(upload_path, extension, raw_bytes)
        logger.info(
            "Stored resume at %s (size=%d bytes, pages=%d)",
            stored_path,
            len(raw_bytes),
            parsed_json["page_count"],
        )
    else:
        logger.debug(
            "Skipping local storage of resume (STORE_LOCAL_UPLOAD_COPY disabled)"
        )
    return ResumeIngestResult(
        file_path=stored_path, parsed_json=parsed_json, warnings=warnings
    )


def _detect_extension(filename: str) -> str:
    """从文件名获取后缀，小写化后返回。"""
    _, ext = os.path.splitext(filename.lower())
    return ext


def _detect_file_type_from_content(payload: bytes) -> Optional[str]:
    """从文件内容检测实际文件类型。"""
    if not payload:
        return None

    # Check file signatures (magic numbers)
    if payload.startswith(b"%PDF-"):
        return ".pdf"
    elif payload.startswith(b"PK\x03\x04"):  # ZIP-based formats
        # Check if it's a DOCX by looking for specific files inside
        try:
            import zipfile

            with zipfile.ZipFile(io.BytesIO(payload)) as zip_file:
                if "word/document.xml" in zip_file.namelist():
                    return ".docx"
                elif "xl/workbook.xml" in zip_file.namelist():
                    return ".xlsx"
        except:
            pass
        return ".zip"  # Generic ZIP file
    elif payload.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):  # OLE2 format
        return ".doc"  # Legacy Word document
    # Removed duplicate ZIP signature branch

    # Check for text content
    try:
        text = payload.decode("utf-8", errors="ignore")
        if len(text.strip()) > 10:  # Has substantial text content
            return ".txt"
    except:
        pass

    return None


def _parse_by_extension(extension: str, payload: bytes) -> Tuple[List[str], List[str]]:
    """按照扩展名分发到具体解析函数。"""
    logger.debug(
        f"Parsing file extension={extension}, payload_size={len(payload)} bytes"
    )

    if extension == ".pdf":
        logger.debug("Using PDF parser")
        return _parse_pdf(payload)
    if extension == ".docx":
        logger.debug("Using DOCX parser")
        return _parse_docx(payload)
    if extension == ".txt":
        logger.debug("Using TXT parser")
        return _parse_txt(payload)
    raise ResumeParsingError(f"Unsupported file type: {extension}")


def _parse_pdf(payload: bytes) -> Tuple[List[str], List[str]]:
    """使用 pypdf 抽取每页文本，并生成对应的警告信息。"""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise ResumeParsingError("pypdf is not installed.") from exc

    warnings: List[str] = []
    pages: List[str] = []
    reader = PdfReader(io.BytesIO(payload))

    logger.debug(f"PDF pages={len(reader.pages)}")

    for index, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            logger.debug(f"Page {index + 1} raw_text_len={len(text)}")
            logger.debug(f"Page {index + 1} preview: {text[:200]}...")
        except Exception as exc:  # pragma: no cover - defensive
            warnings.append(f"Failed to extract text from page {index + 1}: {exc}")
            text = ""
        clean_text = _normalise_text(text)
        logger.debug(f"Page {index + 1} cleaned_len={len(clean_text)}")
        if not clean_text.strip():
            warnings.append(f"Page {index + 1} appears empty.")
        pages.append(clean_text)
    if not pages:
        raise ResumeParsingError("No text extracted from PDF.")
    return pages, warnings


def _parse_docx(payload: bytes) -> Tuple[List[str], List[str]]:
    """使用 python-docx 抽取段落与表格文本。"""
    try:
        import docx  # type: ignore
        import zipfile

        # Check if the file is a valid ZIP (DOCX requirement)
        try:
            with zipfile.ZipFile(io.BytesIO(payload)) as zip_file:
                # Check for required DOCX files
                required_files = ["[Content_Types].xml", "word/document.xml"]
                if not all(f in zip_file.namelist() for f in required_files):
                    raise ResumeParsingError(
                        "File appears to be corrupted or not a valid DOCX file"
                    )
        except zipfile.BadZipFile as e:
            # Provide concise error and move details to debug
            logger.error(f"DOCX ZIP corrupted: {str(e)}")
            logger.debug(f"Corrupted DOCX size={len(payload)} bytes")
            logger.debug(f"Corrupted DOCX signature={payload[:20].hex()}")

            # Try to provide a more user-friendly error message
            if "Bad magic number for central directory" in str(e):
                # Try to extract text using alternative methods for corrupted files
                logger.warning(
                    "Attempting recovery from corrupted DOCX using alternative extraction"
                )
                try:
                    # Try to extract text using python-docx even with corrupted ZIP
                    document = docx.Document(io.BytesIO(payload))
                    blocks: List[str] = []
                    for paragraph in document.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            blocks.append(text)
                    for table in getattr(document, "tables", []):
                        blocks.extend(_flatten_table(table))
                    text = "\n".join(blocks)
                    normalised = _normalise_text(text)
                    if normalised.strip():
                        logger.debug(
                            f"Recovered {len(normalised)} characters from corrupted DOCX"
                        )
                        return [normalised], [
                            "File was corrupted but text extraction succeeded"
                        ]
                except Exception as recovery_error:
                    logger.warning(f"Recovery attempt failed: {recovery_error}")

                raise ResumeParsingError(
                    "The uploaded file appears to be corrupted. Please try re-saving the document and uploading again. "
                    "If the problem persists, try saving as a different format (PDF or plain text)."
                )
            else:
                raise ResumeParsingError(f"File is not a valid DOCX format: {str(e)}")
        except Exception as e:
            logger.error(f"DOCX validation error: {str(e)}")
            raise ResumeParsingError(f"File validation failed: {str(e)}")
    except ImportError as exc:
        raise ResumeParsingError("python-docx is not installed.") from exc

    document = docx.Document(io.BytesIO(payload))
    blocks: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in getattr(document, "tables", []):
        blocks.extend(_flatten_table(table))
    text = "\n".join(blocks)
    normalised = _normalise_text(text)
    if not normalised.strip():
        raise ResumeParsingError("No text extracted from DOCX.")
    return [normalised], []


def _flatten_table(table: Any) -> Iterable[str]:
    """把表格内容按行拉平成文本，避免遗漏技能或项目描述。"""
    for row in getattr(table, "rows", []):
        cells = [
            cell.text.strip() for cell in getattr(row, "cells", []) if cell.text.strip()
        ]
        if cells:
            yield " | ".join(cells)


def _parse_txt(payload: bytes) -> Tuple[List[str], List[str]]:
    """处理纯文本文件，自动检测编码并归一化。"""
    warnings: List[str] = []
    text = _decode_text(payload, warnings)
    normalised = _normalise_text(text)
    if not normalised.strip():
        raise ResumeParsingError("No text extracted from TXT.")
    return [normalised], warnings


def _decode_text(payload: bytes, warnings: List[str]) -> str:
    """尝试常见编码顺序，必要时记录降级警告。"""
    encodings = ["utf-8-sig", "utf-8", "latin-1"]
    for encoding in encodings:
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    warnings.append("Fell back to latin-1 decoding.")
    return payload.decode("latin-1", errors="replace")


def _normalise_text(text: str) -> str:
    """统一字符、换行与空白，确保下游 chunk/LLM 输入干净。"""
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = _normalise_bullets(text)
    text = _collapse_blank_lines(text)
    return text.strip()


def _normalise_bullets(text: str) -> str:
    """将多种项目符号转换成标准短横线前缀。"""
    bullet_chars = {"•", "◦", "▪", "‣"}
    lines = []
    for line in text.split("\n"):
        stripped = line.lstrip()
        if stripped and stripped[0] in bullet_chars:
            line = "- " + stripped[1:].lstrip()
        lines.append(line)
    return "\n".join(lines)


def _collapse_blank_lines(text: str) -> str:
    """限制连续空行数量，提升可读性。"""
    return re.sub(r"\n{3,}", "\n\n", text)


def _store_file(directory: Path, extension: str, payload: bytes) -> Path:
    """以 UUID 命名保存原始文件，返回最终路径。"""
    filename = f"{uuid.uuid4().hex}{extension}"
    path = directory / filename
    path.write_bytes(payload)
    return path


def _build_parsed_payload(
    pages: Sequence[str], warnings: Sequence[str]
) -> Dict[str, Any]:
    """构造存入数据库的标准化 JSON 结构。"""
    combined_text = "\n\n".join(pages).strip()
    words = combined_text.split()
    return {
        "raw_text": combined_text,
        "pages": list(pages),
        "page_count": len(pages),
        "word_count": len(words),
        "warnings": list(warnings),
    }
